from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def export_rewrite_docx(output_path: Path, document: dict, rewrite: dict, benefit_text: Optional[str] = None) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(document.get("title") or Path(document["filename"]).stem)
    run.bold = True
    run.font.name = "宋体"
    run.font.size = Pt(16)

    add_heading(doc, "引言", 1)
    add_body(doc, rewrite.get("intro") or "本研究围绕项目应用场景、关键技术路径和实施效果开展分析，形成面向科技项目验收的研究报告优化稿。")

    add_heading(doc, "1 研究背景与问题分析", 1)
    add_body(doc, rewrite.get("background") or "结合原始材料，项目需要进一步突出研究对象、业务痛点、技术约束和研究意义，避免停留在产品功能或专利说明层面。")

    add_heading(doc, "2 研究内容与技术路线", 1)
    for item in rewrite.get("outline", [])[:6]:
        add_heading(doc, item, 2)
        add_body(doc, "本节建议补充研究方法、技术路线、实施过程、验证指标和数据来源，使报告符合技术论文体例。")

    add_heading(doc, "3 研究成果与应用验证", 1)
    for section in rewrite.get("sections", []):
        add_heading(doc, section.get("title", "优化章节"), 2)
        add_body(doc, section.get("after", ""))

    if benefit_text:
        add_heading(doc, "4 效益分析", 1)
        add_body(doc, benefit_text)

    add_heading(doc, "5 结论", 1)
    add_body(doc, rewrite.get("conclusion") or "本研究形成了可用于项目验收的优化报告框架。后续应补充实测数据、运行周期和第三方证明材料，以增强研究结论的客观性和可复核性。")

    add_heading(doc, "参考文献", 1)
    add_body(doc, "参考文献需按 GB/T 7714 格式补充。")
    doc.save(output_path)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(10)
    fmt.space_after = Pt(10)
    fmt.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "宋体"
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 10.5)


def add_body(doc: Document, text: str) -> None:
    for part in str(text or "").split("\n"):
        if not part.strip():
            continue
        paragraph = doc.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.5
        fmt.first_line_indent = Pt(21)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(part.strip())
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
